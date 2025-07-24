# --- app.py ---
# This is the main Streamlit application file.
# It orchestrates UI, file handling, and calls to LLM agents.

import streamlit as st
import io
import os
import json
import pandas as pd # Import pandas for DataFrame display
import plotly.express as px # For graphs/diagrams
import time # Import the time module for timing operations
from docx import Document # Import Document from python-docx for writing DOCX
from docx.shared import Inches # For potentially adding images, though not used here

# --- IMPORTANT: Load environment variables from .env file ---
# Ensure python-dotenv is installed (pip install python-dotenv)
# and your .env file is in the project root.
from dotenv import load_dotenv
load_dotenv() # This line loads the variables from .env

# Import helper functions and LLM agents
import utils
import agents.llm_connector as llm_connector
# Explicitly import the agent functions
from agents.data_extraction_agent import data_extraction_agent_run
from agents.analysis_agent import analysis_agent_run
from agents.proposal_generation_agent import proposal_generation_agent_run


# --- Streamlit App Configuration (MUST BE THE FIRST STREAMLIT COMMANDS) ---
st.set_page_config(layout="wide", page_title="SOW-to-Proposal AI Assistant", initial_sidebar_state="expanded")

# --- Custom CSS for professional look and responsiveness ---
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;600;700&display=swap');

    html, body, [class*="st-emotion"] {
        font-family: 'Inter', sans-serif;
        color: #2c3e50; /* Darker text for professionalism */
    }

    .stApp {
        background-color: #f8f9fa; /* Light background */
        padding: 1rem 5%; /* Responsive padding */
    }

    /* Adjust main content area padding for better responsiveness */
    .st-emotion-cache-1cypcdb { /* This specific class might change with Streamlit updates */
        padding-top: 2rem;
        padding-bottom: 2rem;
        padding-left: 0; /* Handled by .stApp padding */
        padding-right: 0; /* Handled by .stApp padding */
    }

    h1 {
        color: #1a2a3a;
        text-align: center;
        font-weight: 700;
        margin-bottom: 0.5em;
        font-size: 2.5em; /* Larger for impact */
    }

    h2 {
        color: #34495e;
        border-bottom: 2px solid #e9ecef;
        padding-bottom: 0.5em;
        margin-top: 2em;
        font-weight: 600;
        font-size: 1.8em;
    }

    h3 {
        color: #495057;
        margin-top: 1.5em;
        font-weight: 600;
        font-size: 1.4em;
    }

    /* General button styling */
    .stButton > button {
        background-color: #007bff; /* Primary blue for buttons */
        color: white;
        border-radius: 8px;
        border: none;
        padding: 10px 20px;
        font-size: 16px;
        font-weight: 500;
        transition: all 0.2s ease-in-out;
        box-shadow: 0 2px 5px rgba(0, 123, 255, 0.2);
        width: auto; /* Allow buttons to size naturally */
        display: inline-flex; /* For icon alignment */
        align-items: center;
        justify-content: center;
        gap: 8px;
    }
    .stButton > button:hover {
        background-color: #0056b3;
        box-shadow: 0 4px 10px rgba(0, 123, 255, 0.3);
        transform: translateY(-2px);
    }
    .stButton > button:active {
        transform: translateY(0);
        box-shadow: 0 1px 3px rgba(0, 123, 255, 0.2);
    }

    /* Metric cards */
    div[data-testid="stMetric"] {
        background-color: #ffffff;
        border-radius: 10px;
        padding: 15px;
        box-shadow: 0 2px 8px rgba(0, 0, 0, 0.1);
        border-left: 5px solid #28a745; /* Green accent */
        margin-bottom: 15px; /* Spacing between metrics */
    }
    div[data-testid="stMetricLabel"] {
        font-size: 1.1em;
        font-weight: 600;
        color: #495057;
    }
    div[data-testid="stMetricValue"] {
        font-size: 2em;
        font-weight: 700;
        color: #007bff;
    }

    /* Tabs styling */
    .stTabs [data-testid="stTab"] {
        font-size: 1.1em;
        font-weight: 600;
        color: #6c757d;
        padding: 10px 15px;
        border-radius: 8px 8px 0 0;
        transition: all 0.2s ease-in-out;
    }
    .stTabs [data-testid="stTab"] p {
        font-size: 1.1em;
        margin: 0;
    }
    .stTabs [data-testid="stTab"][aria-selected="true"] {
        color: #007bff;
        border-bottom: 3px solid #007bff;
        background-color: #e9f5ff; /* Light blue background for active tab */
    }
    .stTabs [data-testid="stTab"]:hover {
        background-color: #f0f8ff;
    }

    /* Spinner customization */
    .stSpinner > div {
        color: #007bff;
        font-size: 1.2em;
        /* Ensure spinner text is horizontal */
        white-space: nowrap; /* Prevent text from wrapping */
        display: flex; /* Use flex to align spinner and text */
        align-items: center; /* Vertically align them */
        gap: 10px; /* Space between spinner icon and text */
    }
    .stSpinner > div > div {
        border-width: 4px;
        border-top-color: #007bff;
        border-left-color: #007bff;
        width: 30px;
        height: 30px;
    }

    /* Alert messages */
    .stAlert {
        border-radius: 8px;
        font-size: 0.95em;
    }
    .stAlert.info {
        background-color: #e7f5ff;
        color: #004085;
        border-color: #b8daff;
    }
    .stAlert.success {
        background-color: #d4edda;
        color: #155724;
        border-color: #c3e6cb;
    }
    .stAlert.error {
        background-color: #f8d7da;
        color: #721c24;
        border-color: #f5c6cb;
    }
    .stAlert.warning {
        background-color: #fff3cd;
        color: #856404;
        border-color: #ffeeba;
    }

    /* Custom styling for lists in dashboard */
    .dashboard-list ul {
        list-style-type: none;
        padding-left: 0;
    }
    .dashboard-list li {
        background-color: #e9ecef;
        margin-bottom: 5px;
        padding: 10px 15px;
        border-radius: 6px;
        font-size: 0.95em;
        color: #34495e;
        box-shadow: 0 1px 3px rgba(0,0,0,0.05);
    }

    /* Styling for preformatted text (summary, proposal) */
    pre {
        background-color: #f1f3f5;
        padding: 15px;
        border-radius: 8px;
        overflow-x: auto;
        white-space: pre-wrap;
        word-wrap: break-word;
        border: 1px solid #dee2e6;
        font-size: 0.9em;
        line-height: 1.5;
        color: #333;
    }

    /* Dataframe styling */
    .stDataFrame {
        border-radius: 8px;
        overflow: hidden;
        box-shadow: 0 2px 8px rgba(0, 0, 0, 0.08);
    }
    .stDataFrame table {
        border-collapse: collapse;
        width: 100%;
    }
    .stDataFrame th {
        background-color: #007bff;
        color: white;
        padding: 12px 15px;
        text-align: left;
        font-weight: 600;
    }
    .stDataFrame td {
        padding: 10px 15px;
        border-bottom: 1px solid #dee2e6;
    }
    .stDataFrame tr:nth-child(even) {
        background-color: #f8f9fa;
    }

    /* Responsive adjustments */
    @media (max-width: 768px) {
        .stApp {
            padding: 1rem 2%;
        }
        h1 {
            font-size: 2em;
        }
        h2 {
            font-size: 1.5em;
        }
        .stButton > button {
            width: 100%; /* Full width buttons on small screens */
            margin-bottom: 10px;
        }
        div[data-testid="stMetric"] {
            padding: 10px;
        }
        div[data-testid="stMetricValue"] {
            font-size: 1.5em;
        }
    }
</style>
""", unsafe_allow_html=True)


st.title("📄 SOW-to-Proposal AI Assistant (MVP)")
st.markdown("Upload your Scope of Work (SOW) document to get a detailed summary and a draft technical proposal.")


# --- LLM Choice (Configured in Code) ---
# Choose your LLM provider: "openai" or "gemini"
# Make sure the corresponding API key is set in your .env file
LLM_CHOICE = "gemini" # <--- IMPORTANT: SET YOUR LLM CHOICE HERE! (e.g., "openai" or "gemini")

# Display which LLM is being used in the main content area
st.info(f"💡 Currently using **{LLM_CHOICE.capitalize()}** for AI generation.")

# Initialize LLM clients via the connector
try:
    llm_connector.initialize_llm_clients(LLM_CHOICE)
except ValueError as e:
    st.error(f"LLM Initialization Error: {e}")
    st.stop() # Stop execution if LLM client cannot be initialized


# Function to clear session state for a fresh start
def reset_app():
    for key in st.session_state.keys():
        del st.session_state[key]
    st.rerun()

st.sidebar.button("🔄 Start Fresh / Reset", on_click=reset_app)


# --- File Uploader Section ---
st.header("1. Upload Your Scope of Work (SOW) �")
uploaded_file = st.file_uploader("Select a PDF or DOCX file", type=["pdf", "docx"])

sow_text = ""
if uploaded_file is not None:
    file_extension = uploaded_file.name.split(".")[-1].lower()
    file_buffer = io.BytesIO(uploaded_file.getvalue())

    start_time_extraction = time.time() # Start timer for extraction
    with st.spinner("🚀 Extracting text from SOW... This may take a moment for larger files."):
        if file_extension == "pdf":
            sow_text = utils.extract_text_from_pdf(file_buffer)
        elif file_extension == "docx":
            sow_text = utils.extract_text_from_docx(file_buffer)
        else:
            st.error("Unsupported file type. Please upload a PDF or DOCX.")
            sow_text = ""
    end_time_extraction = time.time() # End timer for extraction

    if sow_text:
        st.success(f"Text extracted successfully! 🎉 (Took {end_time_extraction - start_time_extraction:.2f} seconds)")
        with st.expander("View Raw Extracted Text (Click to expand)"):
            st.text_area("Raw SOW Text", sow_text, height=300, disabled=True)
    else:
        st.error("Could not extract text from the uploaded SOW document. Please ensure it's a readable PDF or DOCX.")

# --- SOW Analysis & Summary Section ---
if sow_text:
    st.header("2. Analyze SOW & Get Summary")

    # Button to trigger SOW analysis
    if st.button("Analyze SOW Details", key="analyze_sow_btn"):
        if not sow_text:
            st.warning("Please upload an SOW document first.")
        else:
            # Check if the selected LLM client is initialized
            if (LLM_CHOICE == "openai" and not llm_connector.openai_client) or \
               (LLM_CHOICE == "gemini" and not llm_connector.gemini_model):
                st.error(f"{LLM_CHOICE.capitalize()} LLM client not properly initialized. Please check your API key setup in .env.")
            else:
                start_time_data_extraction_agent = time.time() # Start timer for Data Extraction Agent
                with st.spinner(" **Data Extraction Agent** is analyzing SOW details..."):
                    # Corrected function call: directly call the imported function
                    sow_structured_data = data_extraction_agent_run(sow_text, llm_choice=LLM_CHOICE)
                    if sow_structured_data and not sow_structured_data.get("error"):
                        st.session_state['sow_structured_data'] = sow_structured_data
                        st.success("Data Extraction Agent complete! ✨")
                    else:
                        st.error(f"Data Extraction Agent failed: {sow_structured_data.get('error', 'Unknown error.')}")
                        st.session_state['sow_structured_data'] = None # Clear state on error
                end_time_data_extraction_agent = time.time() # End timer
                if 'sow_structured_data' in st.session_state and not st.session_state['sow_structured_data'].get("error"):
                    st.info(f"**Data Extraction Agent** took {end_time_data_extraction_agent - start_time_data_extraction_agent:.2f} seconds.")


    # Display SOW Analysis Dashboard
    if 'sow_structured_data' in st.session_state and st.session_state['sow_structured_data'] and \
       not st.session_state['sow_structured_data'].get("error"):
        sow_data = st.session_state['sow_structured_data']

        st.markdown("### 📊 Structured SOW Insights Dashboard")
        st.info("Here's a quick overview of the key information extracted from your SOW, organized for clarity.")

        # --- Dashboard Tabs ---
        tab1, tab2, tab3, tab4, tab5 = st.tabs(["Overview", "Scope", "Deliverables", "Technical", "Constraints & Stakeholders"])

        with tab1:
            st.subheader("Project Overview")
            col_proj_name, col_client_name, col_timeline = st.columns(3)
            with col_proj_name:
                st.metric("Project Name", sow_data.get('project_name', 'N/A'))
            with col_client_name:
                st.metric("Client Name", sow_data.get('client_name', 'N/A'))
            with col_timeline:
                # Display timeline as a joined string for metric, or list in detail section
                timeline_display = " | ".join(sow_data.get('timeline_overview', ['N/A'])) if sow_data.get('timeline_overview') else 'N/A'
                st.metric("Timeline Overview", timeline_display)

            st.markdown("---")
            st.markdown("#### Key Objectives")
            if sow_data.get("objectives"):
                st.markdown('<div class="dashboard-list"><ul>', unsafe_allow_html=True)
                for i, obj in enumerate(sow_data["objectives"]):
                    st.markdown(f"<li>**{i+1}.** {obj}</li>", unsafe_allow_html=True)
                st.markdown('</ul></div>', unsafe_allow_html=True)
            else:
                st.info("No specific objectives identified in the SOW.")
            
            # Conditional Chart: Bar chart for count of objectives
            if sow_data.get("objectives") and len(sow_data["objectives"]) > 0:
                obj_count = len(sow_data["objectives"])
                obj_count_df = pd.DataFrame({'Category': ['Key Objectives'], 'Count': [obj_count]})
                
                # Make chart more flexible: only show if count > 0, and auto-scale y-axis
                fig = px.bar(obj_count_df, x='Category', y='Count', title='Number of Key Objectives Identified', text_auto=True,
                             color_discrete_sequence=['#28a745']) # Green color
                fig.update_traces(texttemplate='%{text}', textposition='outside')
                fig.update_layout(xaxis_title="", yaxis_title="Count", showlegend=False,
                                  yaxis_range=[0, obj_count + 1] if obj_count < 5 else None, # Auto-scale for larger counts
                                  height=300) # Fixed height for consistency
                st.plotly_chart(fig, use_container_width=True)
            else:
                st.markdown("*(Chart not available: No objectives data to plot.)*")


        with tab2:
            st.subheader("Scope Definition")
            col_scope_in, col_scope_out = st.columns(2)
            with col_scope_in:
                st.markdown("**Included Scope of Work:**")
                if sow_data.get("scope_of_work"):
                    st.markdown('<div class="dashboard-list"><ul>', unsafe_allow_html=True)
                    for i, item in enumerate(sow_data["scope_of_work"]):
                        st.markdown(f"<li>**{i+1}.** {item}</li>", unsafe_allow_html=True)
                    st.markdown('</ul></div>', unsafe_allow_html=True)
                else:
                    st.info("No detailed included scope identified.")
            with col_scope_out:
                st.markdown("**Out of Scope:**")
                if sow_data.get("out_of_scope"):
                    st.markdown('<div class="dashboard-list"><ul>', unsafe_allow_html=True)
                    for i, item in enumerate(sow_data["out_of_scope"]):
                        st.markdown(f"<li>**{i+1}.** {item}</li>", unsafe_allow_html=True)
                    st.markdown('</ul></div>', unsafe_allow_html=True)
                else:
                    st.info("No specific out-of-scope items identified.")

            # Conditional Chart: Pie chart for scope distribution
            scope_counts = {'Included Scope': len(sow_data.get("scope_of_work", [])),
                            'Out of Scope': len(sow_data.get("out_of_scope", []))}
            
            if sum(scope_counts.values()) > 0: # Only show chart if there's any scope data
                scope_df = pd.DataFrame(list(scope_counts.items()), columns=['Category', 'Count'])
                fig_pie = px.pie(scope_df, values='Count', names='Category', title='Scope Items Distribution',
                                 color_discrete_sequence=px.colors.sequential.RdBu,
                                 height=350) # Fixed height
                st.plotly_chart(fig_pie, use_container_width=True)
            else:
                st.markdown("*(Chart not available: No scope data to plot.)*")


        with tab3:
            st.subheader("Key Deliverables")
            if sow_data.get("deliverables"):
                # Ensure deliverables are in a suitable format for DataFrame
                cleaned_deliverables = []
                for item in sow_data["deliverables"]:
                    if isinstance(item, dict) and "name" in item and "description" in item:
                        cleaned_deliverables.append(item)
                    elif isinstance(item, str): # Handle if LLM returned list of strings
                         cleaned_deliverables.append({"name": item, "description": "Not detailed by AI"})
                
                if cleaned_deliverables:
                    deliverables_df = pd.DataFrame(cleaned_deliverables)
                    st.dataframe(deliverables_df, use_container_width=True, hide_index=True)
                else:
                    st.info("No specific deliverables identified or could not be parsed into a table.")
            else:
                st.info("No specific deliverables identified.")

        with tab4:
            st.subheader("Technical Landscape")
            st.markdown("**Technical Requirements:**")
            if sow_data.get("technical_requirements"):
                st.markdown('<div class="dashboard-list"><ul>', unsafe_allow_html=True)
                for i, tech in enumerate(sow_data["technical_requirements"]):
                    st.markdown(f"<li>**{i+1}.** {tech}</li>", unsafe_allow_html=True)
                st.markdown('</ul></div>', unsafe_allow_html=True)
            else:
                st.info("No specific technical requirements identified.")
            
            # Conditional Chart: Bar chart for number of technical requirements
            if sow_data.get("technical_requirements") and len(sow_data["technical_requirements"]) > 0:
                tech_count = len(sow_data["technical_requirements"])
                tech_count_df = pd.DataFrame({'Category': ['Technical Requirements'], 'Count': [tech_count]})
                fig = px.bar(tech_count_df, x='Category', y='Count', title='Number of Technical Requirements Identified', text_auto=True,
                             color_discrete_sequence=['#17a2b8']) # Info color
                fig.update_traces(texttemplate='%{text}', textposition='outside')
                fig.update_layout(xaxis_title="", yaxis_title="Count", showlegend=False,
                                  yaxis_range=[0, tech_count + 1] if tech_count < 5 else None, # Auto-scale for smaller counts
                                  height=300) # Fixed height
                st.plotly_chart(fig, use_container_width=True)
            else:
                st.markdown("*(Chart not available: No technical requirements data to plot.)*")

        with tab5:
            st.subheader("Constraints & Key Stakeholders")
            col_constr, col_stake_detail = st.columns(2)
            with col_constr:
                st.markdown("**Key Constraints:**")
                if sow_data.get("key_constraints"):
                    st.markdown('<div class="dashboard-list"><ul>', unsafe_allow_html=True)
                    for i, constr in enumerate(sow_data["key_constraints"]):
                        st.markdown(f"<li>**{i+1}.** {constr}</li>", unsafe_allow_html=True)
                    st.markdown('</ul></div>', unsafe_allow_html=True)
                else:
                    st.info("No specific constraints identified.")
            with col_stake_detail:
                st.markdown("**Key Stakeholders:**")
                if sow_data.get("stakeholders"):
                    st.markdown('<div class="dashboard-list"><ul>', unsafe_allow_html=True)
                    for i, stakeholder in enumerate(sow_data["stakeholders"]):
                        st.markdown(f"<li>**{i+1}.** {stakeholder}</li>", unsafe_allow_html=True)
                    st.markdown('</ul></div>', unsafe_allow_html=True)
                else:
                    st.info("No specific stakeholders identified.")

        st.markdown("---")
        st.markdown("### 📝 Detailed Narrative Summary of SOW")
        start_time_summary = time.time() # Start timer for summary
        with st.spinner("✍️ **Analysis Agent** is generating detailed narrative summary..."):
            # Corrected function call: directly call the imported function
            detailed_summary = analysis_agent_run(sow_data, llm_choice=LLM_CHOICE)
            st.write(detailed_summary)
        end_time_summary = time.time() # End timer for summary
        st.info(f"**Analysis Agent** (Summary Generation) took {end_time_summary - start_time_summary:.2f} seconds.")
        st.markdown("---")

        # --- Draft Technical Proposal Section ---
        st.header("3. Generate Technical Proposal Draft 🚀")
        st.info("This will generate a preliminary technical proposal based on the SOW analysis. Remember, this is a draft and requires human review and refinement.")
        if st.button("Generate Proposal Draft", key="generate_proposal_btn"):
            # Check if the selected LLM client is initialized
            if (LLM_CHOICE == "openai" and not llm_connector.openai_client) or \
               (LLM_CHOICE == "gemini" and not llm_connector.gemini_model):
                st.error(f"{LLM_CHOICE.capitalize()} LLM client not properly initialized. Please check your API key setup in .env.")
            else:
                start_time_proposal = time.time() # Start timer for proposal
                with st.spinner("🤖 **Proposal Generation Agent** is drafting technical proposal... This will take a moment."):
                    # Corrected function call: directly call the imported function
                    proposal_draft = proposal_generation_agent_run(sow_data, llm_choice=LLM_CHOICE)
                    st.session_state['proposal_draft'] = proposal_draft
                    st.success("Technical proposal draft generated!")
                end_time_proposal = time.time() # End timer for proposal
                st.info(f"**Proposal Generation Agent** took {end_time_proposal - start_time_proposal:.2f} seconds.")

        if 'proposal_draft' in st.session_state and st.session_state['proposal_draft']:
            st.markdown("### ✍️ Your Generated Proposal Draft")
            st.markdown(st.session_state['proposal_draft'])

            # Buttons for download and copy
            col_dl_md, col_dl_docx, col_copy = st.columns(3) # Added a column for DOCX download
            with col_dl_md:
                st.download_button(
                    label="Download as Markdown",
                    data=st.session_state['proposal_draft'],
                    file_name="technical_proposal_draft.md",
                    mime="text/markdown"
                )
            with col_dl_docx:
                # Generate DOCX content
                docx_buffer = io.BytesIO()
                doc = Document()
                # Simple conversion from Markdown to DOCX (basic headers and paragraphs)
                for line in st.session_state['proposal_draft'].split('\n'):
                    if line.startswith('## '):
                        doc.add_heading(line.replace('## ', ''), level=2)
                    elif line.startswith('# '):
                        doc.add_heading(line.replace('# ', ''), level=1)
                    elif line.startswith('- '):
                        doc.add_paragraph(line.replace('- ', ''), style='List Bullet')
                    else:
                        doc.add_paragraph(line)
                doc.save(docx_buffer)
                docx_buffer.seek(0) # Rewind the buffer to the beginning

                st.download_button(
                    label="Download as DOCX",
                    data=docx_buffer.getvalue(),
                    file_name="technical_proposal_draft.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            with col_copy:
                st.text_area("Copy Proposal Draft (Manual Copy)", st.session_state['proposal_draft'], height=200, help="Select all text and copy manually.")
                st.info("For security reasons, direct 'Copy to Clipboard' is not available in Streamlit. Please manually select and copy the text from the box above.")

# --- Footer ---
st.markdown("---")
st.markdown("Built for Masters Project")